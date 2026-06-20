# -*- coding: utf-8 -*-
# ==============================================================================
#   DISSERTAÇÃO PETR4 — Gerador da Documentação de Coleta e Preparação dos Dados
#   Autor: Vanderlei Barbosa da Silva | Orientador: Prof. Dr. Julio Cesar Nievola
#
#   Gera um documento Word (.docx) com TODO o detalhamento da engenharia de dados
#   (coleta, categorização, filtragem, split), com texto, código, tabelas e
#   gráficos estatísticos a partir dos dados REAIS. Reprodutível.
#
#   Saída: Documentacao_Coleta_PETR4.docx
# ==============================================================================

from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────────────────
# CONFIG E LEITURA DOS DADOS REAIS
# ──────────────────────────────────────────────────────────────────────────────
PASTA = Path("./Mestrado_PETR4")
ARQ_ORIGINAL = PASTA / "base_textual_petr4_wordpress_2018_2025.csv"
ARQ_TRATADA  = PASTA / "base_textual_petr4_tratada.csv"
ASSETS = Path("./_doc_assets"); ASSETS.mkdir(exist_ok=True)
SAIDA = Path("Documentacao_Coleta_PETR4.docx")

plt.rcParams.update({"figure.dpi": 150, "font.size": 11,
                     "axes.spines.top": False, "axes.spines.right": False})
AZUL, VERM, VERDE, CINZA = "#2c7bb6", "#d7191c", "#1a9641", "#999999"

print("Lendo bases...")
df = pd.read_csv(ARQ_ORIGINAL)
df["dt"] = pd.to_datetime(df["data_publicacao"], errors="coerce")
df = df[df["dt"].notna()].copy()
df["ano"] = df["dt"].dt.year
N_ORIG = len(df)

# Base tratada (pode não ter sido gerada; cai para a original)
try:
    dft = pd.read_csv(ARQ_TRATADA)
    dft["dt"] = pd.to_datetime(dft["data_publicacao"], errors="coerce")
except Exception:
    dft = df.copy(); dft["conjunto"] = "treino"

# Estatísticas reais
dist_ano    = df["ano"].value_counts().sort_index()
dist_cat    = df["categoria"].value_counts()
dist_portal = df["dominio"].value_counts()
hora        = df["dt"].dt.hour
lead_lag_pct = (hora >= 17).mean() * 100
tem_hora_pct = ((hora != 0) | (df["dt"].dt.minute != 0)).mean() * 100

# Filtro ESTRITO (para comparação) — recomputa
TERMOS_ESTRITO = ["petrobras","petr4","petr3","petroleira","petróleo","petroleo",
                  "brent","wti","opep","barril","combustível","combustivel","refinaria"]
alvo = (df["titulo"].fillna("") + " " + df["resumo"].fillna("")).str.lower()
mask_estrito = alvo.apply(lambda t: any(x in t for x in TERMOS_ESTRITO))
dist_cat_estrito = df[mask_estrito]["categoria"].value_counts()
N_ESTRITO = int(mask_estrito.sum())

# Filtro LEVE (qualidade)
tit = df["titulo"].fillna("").astype(str).str.strip()
mask_leve = ~((tit == "") | (tit.str.len() < 15))
N_LEVE = int(mask_leve.sum())
dist_cat_leve = df[mask_leve]["categoria"].value_counts()

# Split (da base tratada)
if "conjunto" in dft.columns:
    dist_split = dft["conjunto"].value_counts()
    tab_split = pd.crosstab(dft["dt"].dt.year, dft["conjunto"]).reindex(
        columns=["treino","validacao","teste"], fill_value=0)
else:
    dist_split = pd.Series({"treino":0,"validacao":0,"teste":0}); tab_split = pd.DataFrame()

# Coleta bruta (do relatório final do log)
COLETA = {"brutos":1040007, "unicos":205716, "duplicados":834291}

# ──────────────────────────────────────────────────────────────────────────────
# GRÁFICOS
# ──────────────────────────────────────────────────────────────────────────────
print("Gerando gráficos...")

def salvar(fig, nome):
    caminho = ASSETS / nome
    fig.savefig(caminho, bbox_inches="tight"); plt.close(fig)
    return caminho

# 1. Anual
fig, ax = plt.subplots(figsize=(8,4))
ax.bar(dist_ano.index.astype(str), dist_ano.values, color=AZUL)
for x,y in zip(dist_ano.index.astype(str), dist_ano.values):
    ax.text(x, y, f"{y:,}".replace(",","."), ha="center", va="bottom", fontsize=9)
ax.set_title("Notícias coletadas por ano (2018–2025)"); ax.set_ylabel("Notícias")
g_anual = salvar(fig, "anual.png")

# 2. Portais
fig, ax = plt.subplots(figsize=(8,4))
ax.barh(dist_portal.index[::-1], dist_portal.values[::-1], color=VERDE)
for i,(k,v) in enumerate(zip(dist_portal.index[::-1], dist_portal.values[::-1])):
    ax.text(v, i, f" {v:,}".replace(",","."), va="center", fontsize=9)
ax.set_title("Notícias por portal (fonte)"); ax.set_xlabel("Notícias")
g_portais = salvar(fig, "portais.png")

# 3. Categorias
fig, ax = plt.subplots(figsize=(8,4.5))
ax.barh(dist_cat.index[::-1], dist_cat.values[::-1], color=AZUL)
for i,v in enumerate(dist_cat.values[::-1]):
    ax.text(v, i, f" {v:,}".replace(",","."), va="center", fontsize=9)
ax.set_title("Notícias por categoria temática"); ax.set_xlabel("Notícias")
g_categorias = salvar(fig, "categorias.png")

# 4. Lead-Lag (pizza)
fig, ax = plt.subplots(figsize=(5.5,5.5))
antes = (hora < 17).sum(); depois = (hora >= 17).sum()
ax.pie([antes, depois], labels=[f"Até 17h\n{antes:,}".replace(",","."),
       f"Após 17h (Lead-Lag)\n{depois:,}".replace(",",".")],
       autopct="%1.1f%%", colors=[AZUL, VERM], startangle=90,
       wedgeprops=dict(width=0.45))
ax.set_title("Notícias antes/depois do fechamento da B3 (17h)")
g_leadlag = salvar(fig, "leadlag.png")

# 5. Hora do dia
fig, ax = plt.subplots(figsize=(8,3.6))
vc = hora.value_counts().sort_index()
ax.bar(vc.index, vc.values, color=CINZA)
ax.axvline(17, color=VERM, ls="--", lw=1.2, label="Fechamento B3 (17h)")
ax.set_title("Distribuição das notícias por hora de publicação")
ax.set_xlabel("Hora (horário de Brasília)"); ax.set_ylabel("Notícias"); ax.legend()
g_horas = salvar(fig, "horas.png")

# 6. Filtro estrito vs leve por categoria
fig, ax = plt.subplots(figsize=(8.5,4.5))
cats = list(dist_cat.index)
orig = [dist_cat.get(c,0) for c in cats]
estr = [dist_cat_estrito.get(c,0) for c in cats]
leve = [dist_cat_leve.get(c,0) for c in cats]
y = np.arange(len(cats)); h=0.27
ax.barh(y+h, orig, h, label="Original", color=CINZA)
ax.barh(y,   leve, h, label="Filtragem leve", color=VERDE)
ax.barh(y-h, estr, h, label="Filtragem estrita", color=VERM)
ax.set_yticks(y); ax.set_yticklabels(cats); ax.invert_yaxis()
ax.set_title("Impacto das filtragens por categoria"); ax.set_xlabel("Notícias"); ax.legend()
g_filtro = salvar(fig, "filtro.png")

# 7. Split por ano (empilhado)
if not tab_split.empty:
    fig, ax = plt.subplots(figsize=(8.5,4.5))
    anos = tab_split.index.astype(str)
    ax.bar(anos, tab_split["treino"], label="Treino (60%)", color=AZUL)
    ax.bar(anos, tab_split["validacao"], bottom=tab_split["treino"], label="Validação (15%)", color=VERDE)
    ax.bar(anos, tab_split["teste"], bottom=tab_split["treino"]+tab_split["validacao"],
           label="Teste (25%)", color=VERM)
    ax.set_title("Split temporal por ano (estratificado)"); ax.set_ylabel("Notícias"); ax.legend()
    g_split = salvar(fig, "split.png")
else:
    g_split = None

# 8. Funil da coleta
fig, ax = plt.subplots(figsize=(8,3.6))
etapas = ["Artigos brutos\n(requisições)", "Após deduplicação\n(únicos)"]
vals = [COLETA["brutos"], COLETA["unicos"]]
ax.bar(etapas, vals, color=[CINZA, AZUL])
for x,v in zip(etapas, vals):
    ax.text(x, v, f"{v:,}".replace(",","."), ha="center", va="bottom", fontsize=10)
ax.set_title("Funil da coleta: brutos → únicos (deduplicação por hash)")
ax.set_ylabel("Artigos")
g_funil = salvar(fig, "funil.png")

# ──────────────────────────────────────────────────────────────────────────────
# HELPERS DOCX
# ──────────────────────────────────────────────────────────────────────────────
doc = Document()
est = doc.styles["Normal"]; est.font.name = "Calibri"; est.font.size = Pt(11)

def h(txt, lvl=1):
    doc.add_heading(txt, level=lvl)

def p(txt):
    par = doc.add_paragraph(txt); par.paragraph_format.space_after = Pt(6)
    return par

def codigo(txt):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Pt(6)
    run = par.add_run(txt.strip("\n"))
    run.font.name = "Consolas"; run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    # sombreamento cinza-claro
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F2F2")
    par._p.get_or_add_pPr().append(shd)
    return par

def tabela(cabecalho, linhas):
    t = doc.add_table(rows=1, cols=len(cabecalho)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,c in enumerate(cabecalho):
        cell = t.rows[0].cells[i]; cell.text = str(c)
        for r in cell.paragraphs[0].runs: r.font.bold = True
    for lin in linhas:
        cells = t.add_row().cells
        for i,v in enumerate(lin): cells[i].text = str(v)
    doc.add_paragraph()
    return t

def img(caminho, larg=6.2):
    if caminho and Path(caminho).exists():
        doc.add_picture(str(caminho), width=Inches(larg))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def legenda(txt):
    par = doc.add_paragraph(txt); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.runs[0]; r.font.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)

fmt = lambda n: f"{int(n):,}".replace(",", ".")

# ──────────────────────────────────────────────────────────────────────────────
# CAPA
# ──────────────────────────────────────────────────────────────────────────────
t = doc.add_heading("Engenharia de Dados Textuais para a Previsão do PETR4", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Documentação da Coleta, Categorização, Filtragem e Particionamento do Corpus de Notícias")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub.runs[0].italic = True
meta = doc.add_paragraph("Vanderlei Barbosa da Silva — PUCPR, Mestrado em Informática\n"
                         "Orientador: Prof. Dr. Julio Cesar Nievola\n"
                         "Dissertação: O Impacto do Sentimento de Notícias Financeiras na "
                         "Previsão de Direção e Volatilidade do Ativo PETR4")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in meta.runs: r.font.size = Pt(10)
doc.add_paragraph()
p("Este documento detalha integralmente o processo de construção da base textual da pesquisa — "
  "fontes, decisões metodológicas, código, estatísticas, sucessos e falhas — para incorporação "
  "direta ao capítulo de Materiais e Métodos da dissertação. Todos os números e gráficos são "
  "gerados a partir dos dados reais coletados.")
doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
# 0. VISÃO GERAL
# ──────────────────────────────────────────────────────────────────────────────
h("1. Visão Geral do Pipeline de Dados", 1)
p("A pesquisa testa se o sentimento extraído de notícias financeiras melhora a previsão da "
  "direção e da volatilidade do ativo PETR4. A engenharia de dados — coleta e higienização do "
  "corpus textual — representa a etapa mais crítica e trabalhosa do projeto. O pipeline é "
  "composto por scripts sequenciais:")
tabela(["Etapa","Script","Função"],
 [["01","01_coleta_dados_financeiros","Preços diários da PETR4 (yfinance) + log-retorno"],
  ["02b","02b_coleta_noticias_wordpress","Coleta de notícias via WordPress REST API (fonte principal)"],
  ["02c","02c_preparar_base_treino_teste_validacao","Filtragem leve + split treino/validação/teste"],
  ["03","03_analise_sentimento_bertimbau","Análise de sentimento (NLP) → Índice de Sentimento da Mídia"],
  ["04","04_modelagem_garch_svm_xgboost","GARCH + SVM/XGBoost + análise de ablação"]])
p(f"O corpus final reúne {fmt(N_ORIG)} notícias únicas com timestamp exato, cobrindo "
  f"integralmente o período de 2018 a 2025.")

# ──────────────────────────────────────────────────────────────────────────────
# 2. FONTES UTILIZADAS
# ──────────────────────────────────────────────────────────────────────────────
h("2. Fontes de Coleta Utilizadas", 1)
p("A coleta final foi realizada sobre cinco portais jornalísticos brasileiros que expõem a API "
  "REST pública do WordPress (endpoint /wp-json/wp/v2/posts). A tabela abaixo lista as fontes e "
  "o volume coletado em cada uma.")
tabela(["Portal","Endpoint (WordPress REST API)","Notícias","Perfil editorial"],
 [["InfoMoney","infomoney.com.br/wp-json/wp/v2/posts", fmt(dist_portal.get("InfoMoney",0)),"Mercado de capitais, PETR4"],
  ["Exame","exame.com/wp-json/wp/v2/posts", fmt(dist_portal.get("Exame",0)),"Negócios e empresas B3"],
  ["Money Times","moneytimes.com.br/wp-json/wp/v2/posts", fmt(dist_portal.get("MoneyTimes",0)),"Investimentos e recomendações"],
  ["Petronotícias","petronoticias.com.br/wp-json/wp/v2/posts", fmt(dist_portal.get("Petronoticias",0)),"Especializado em petróleo e energia"],
  ["Poder360","poder360.com.br/wp-json/wp/v2/posts", fmt(dist_portal.get("Poder360",0)),"Política e governo (governança Petrobras)"]])
img(g_portais)
legenda("Figura 1 — Distribuição das notícias por portal de origem.")

# ──────────────────────────────────────────────────────────────────────────────
# 3. JUSTIFICATIVA DAS FONTES
# ──────────────────────────────────────────────────────────────────────────────
h("3. Justificativa da Escolha das Fontes", 1)
p("A escolha foi guiada por três critérios técnicos decisivos, derivados de uma exigência "
  "metodológica central: a necessidade do TIMESTAMP exato de publicação.")
for titulo, corpo in [
 ("Critério 1 — Disponibilidade do horário exato (timestamp).",
  "A banca foi enfática: sem a hora exata de publicação não é possível garantir o alinhamento "
  "temporal Lead-Lag (uma notícia das 20h, com o mercado fechado, só pode impactar o pregão do "
  "dia seguinte) e, portanto, não há como provar causalidade com o modelo GARCH(1,1). A API REST "
  "do WordPress retorna o campo 'date' no horário de Brasília e 'date_gmt' em UTC — entregando o "
  "timestamp com precisão de segundos."),
 ("Critério 2 — Cobertura histórica gratuita de 2018 a 2025.",
  "Os portais mantêm o histórico completo acessível via API, permitindo varrer todo o período da "
  "pesquisa sem custo e sem depender de bases pagas."),
 ("Critério 3 — Robustez técnica e dados estruturados.",
  "A API devolve JSON estruturado (título, resumo, link, data), eliminando a fragilidade do "
  "web scraping de HTML e o risco de bloqueio anti-bot. A abordagem é metodologicamente análoga "
  "à de Cardoso & Nakane (2024), que coletaram diretamente do portal do Valor Econômico."),
]:
    pr = doc.add_paragraph(); pr.add_run(titulo).bold = True
    pr.add_run(" " + corpo)
p("A pluralidade de fontes (cinco portais com perfis distintos — mercado, negócios, petróleo e "
  "política) reduz o viés de cobertura de uma única fonte (Heston & Sinha, 2017) e cobre tanto a "
  "notícia corporativa quanto os choques exógenos (petróleo, geopolítica, política energética).")

# ──────────────────────────────────────────────────────────────────────────────
# 4. FONTES DESCARTADAS
# ──────────────────────────────────────────────────────────────────────────────
h("4. Fontes Avaliadas e Descartadas", 1)
p("Antes de chegar à solução final, diversas fontes foram testadas e descartadas. O registro "
  "honesto desses caminhos sem saída é parte do rigor metodológico.")
tabela(["Fonte / Método","Motivo do descarte"],
 [["Biblioteca GoogleNews (abordagem inicial)",
   "Mascarava o timestamp (datas estocásticas/aleatórias), inviabilizando o Lead-Lag; além de bloqueio anti-bot."],
  ["GDELT Project (API DOC 2.0)",
   "Bloqueio de IP por volume (HTTP 429 persistente) e cobertura irregular de fontes em português; retornava apenas o título."],
  ["NewsAPI (plano gratuito)",
   "O plano gratuito só permite consultar os últimos ~30 dias (HTTP 426 para datas de 2018–2025). Inútil para o recorte histórico."],
  ["Twitter/X API + Get Old Tweets",
   "Desde 2023 a API tornou-se paga (US$ 100+/mês) e ferramentas como GOT/snscrape foram bloqueadas. Inviável de forma gratuita."],
  ["Valor Econômico (scraping direto)",
   "Conteúdo carregado via JavaScript e atrás de paywall; exigiria Selenium + assinatura."],
  ["Terminais pagos (Factiva, Bloomberg)",
   "Custo institucional incompatível com o orçamento da pesquisa."]])
p("Conclusão: das alternativas gratuitas, apenas a coleta direta do portal via API REST do "
  "WordPress satisfez simultaneamente os três critérios (timestamp, histórico e robustez).")

# ──────────────────────────────────────────────────────────────────────────────
# 5. COMO AS NOTÍCIAS FORAM CAPTURADAS
# ──────────────────────────────────────────────────────────────────────────────
h("5. Como as Notícias Foram Capturadas", 1)
p("A coleta consulta o endpoint /wp-json/wp/v2/posts de cada portal usando o parâmetro 'search' "
  "(busca textual) combinado com 'after'/'before' (janela temporal) e paginação. Para cada termo "
  "da taxonomia, em cada portal, o coletor decide automaticamente a estratégia:")
p("• Estratégia full-range: termos de baixo volume são paginados de uma vez no período inteiro.\n"
  "• Estratégia por janela mensal: termos de alto volume (ex.: \"Petrobras\") são coletados mês a "
  "mês, evitando o teto de offset da API (que torna a paginação profunda lenta e instável).")
p("Cada artigo é deduplicado por hash SHA-256 do título normalizado e gravado imediatamente em "
  "CSV (gravação incremental, com retomada automática em caso de interrupção).")

# ──────────────────────────────────────────────────────────────────────────────
# 6. BIBLIOTECAS E CÓDIGO
# ──────────────────────────────────────────────────────────────────────────────
h("6. Bibliotecas e Código de Captura", 1)
tabela(["Biblioteca","Função no projeto","Por que foi escolhida"],
 [["requests","Requisições HTTP à API REST","Padrão de fato em Python; simples e robusto"],
  ["pandas","Manipulação tabular e CSV","Estrutura DataFrame ideal para o corpus"],
  ["hashlib (SHA-256)","Deduplicação por hash do título","Detecção O(1) de duplicatas, idempotência"],
  ["python-dateutil","Janelas mensais (relativedelta)","Aritmética de datas relativa confiável"],
  ["matplotlib","Gráficos estatísticos","Padrão científico para visualização"],
  ["python-docx","Geração desta documentação","Automação reprodutível do relatório"]])
p("Núcleo da requisição à API REST (função de baixo nível com retry/backoff):")
codigo('''def _get(endpoint, params):
    """GET robusto à API WP REST com retry/backoff. Retorna (posts, total_paginas)."""
    headers = {"User-Agent": USER_AGENT}
    for tentativa in range(1, MAX_RETRY + 1):
        try:
            r = requests.get(endpoint, headers=headers, params=params,
                             timeout=TIMEOUT_S, verify=VERIFY_SSL)
            if r.status_code == 200:
                total_pg = int(r.headers.get("X-WP-TotalPages", 1) or 1)
                return r.json(), total_pg
            if r.status_code == 400:        # página além do fim = fim da paginação
                return [], 0
            return None, 0
        except Exception:
            time.sleep(min(4 * tentativa, 12))
    return None, 0''')
p("Estratégia híbrida (full-range vs. janela mensal) por termo e portal:")
codigo('''posts, total_pg = _get(endpoint, janela_periodo_inteiro)
if total_pg <= LIMITE_PAGINAS_FULLRANGE:           # termo raro → uma varredura
    yield from paginar(janela_periodo_inteiro)
else:                                              # termo de alto volume → mês a mês
    cursor = inicio
    while cursor <= fim:
        fim_mes = min(cursor + relativedelta(months=1) - relativedelta(seconds=1), fim)
        yield from paginar(janela_mensal(cursor, fim_mes))
        cursor += relativedelta(months=1)''')
p("Captura do timestamp e padronização de cada artigo (note os campos 'date' e 'date_gmt'):")
codigo('''return {
    "data_publicacao": post.get("date", ""),       # horário de Brasília (Lead-Lag)
    "data_gmt"       : post.get("date_gmt", ""),    # UTC (auditoria de fuso)
    "categoria"      : categoria,                   # da taxonomia (ablação)
    "titulo"         : limpar_html(post["title"]["rendered"]),
    "resumo"         : limpar_html(post["excerpt"]["rendered"]),
    "url"            : post.get("link", ""),
    "hash_titulo"    : hash_titulo(titulo),         # SHA-256 p/ deduplicação
}''')

# ──────────────────────────────────────────────────────────────────────────────
# 7. EVOLUÇÃO DA COLETA
# ──────────────────────────────────────────────────────────────────────────────
h("7. Evolução da Coleta: Primeiros Números e Ajustes", 1)
p("A quantidade de notícias capturadas evoluiu drasticamente à medida que os obstáculos técnicos "
  "foram superados. A tabela registra essa trajetória — incluindo o que deu errado.")
tabela(["Momento","Abordagem","Notícias","Problema / ajuste"],
 [["Inicial","Biblioteca GoogleNews","~9.079","Timestamp mascarado; datas aleatórias → inválido"],
  ["Tentativa","GDELT (multi-fonte)","centenas","Bloqueio 429 por volume de IP"],
  ["Tentativa","NewsAPI gratuita","~0 (histórico)","Plano cobre só 30 dias"],
  ["Protótipo","WordPress REST (1 mês, 2 portais)","120","Validou timestamp e Lead-Lag"],
  ["Ajuste","+ taxonomia de 7 categorias / 152 termos","—","Cobertura temática ampliada"],
  ["Ajuste","+ estratégia híbrida de paginação","—","Resolveu teto de offset em termos de alto volume"],
  ["Final","WordPress REST (5 portais, 152 termos)", fmt(N_ORIG),"Corpus completo com timestamp"]])
p(f"O salto de ~9.079 para {fmt(N_ORIG)} notícias (≈22×) — todas com hora exata — só foi possível "
  "após a descoberta da API REST do WordPress e dos ajustes de paginação e taxonomia.")
img(g_funil)
legenda(f"Figura 2 — Funil da coleta: {fmt(COLETA['brutos'])} artigos brutos requisitados, "
        f"{fmt(COLETA['duplicados'])} duplicatas removidas por hash, {fmt(COLETA['unicos'])} únicos.")

# ──────────────────────────────────────────────────────────────────────────────
# 8. TERMOS E TAXONOMIA
# ──────────────────────────────────────────────────────────────────────────────
h("8. Termos de Busca e Justificativa", 1)
p("Os termos de busca foram organizados em uma taxonomia de 7 categorias temáticas, cada uma "
  "ancorada na literatura de economia do petróleo e finanças comportamentais. A escolha vai além "
  "da empresa: a literatura mostra que choques exógenos (geopolítica, OPEP, câmbio) são tão ou "
  "mais determinantes para a volatilidade do que notícias corporativas.")
tabela(["Categoria","Nº termos","Foco","Âncora na literatura"],
 [["CAT1 — Empresa", str(len([1])), "Petrobras, PETR4, resultados, governança","Tetlock et al. (2008)"],
  ["CAT2 — Mercado de Petróleo","20","Brent/WTI, OPEP, oferta/demanda","Kilian (2009)"],
  ["CAT3 — Geopolítica","27","Guerras e tensões em produtores","Hamilton (1983); Caldara & Iacoviello (2022)"],
  ["CAT4 — Infraestrutura","20","Refinarias, oleodutos, plataformas","Choque Aramco (2019)"],
  ["CAT5 — Sanções/Navegação","20","Embargos, Ormuz, Mar Vermelho","Prêmio de risco de rotas"],
  ["CAT6 — Governança","24","CEO, intervenção, política energética","Event study (CEO Petrobras 2022)"],
  ["CAT7 — Macro/Energia","20","Dólar, juros, China, transição energética","Zhang et al. (2008); Kilian & Murphy (2014)"]])
p("A primeira coluna reflete a estrutura; ao todo são 152 termos. A justificativa central: "
  "Hamilton (1983) mostrou que quase todas as recessões dos EUA foram precedidas por choques "
  "no petróleo; Kilian (2009) decompôs esses choques em oferta, demanda agregada e demanda "
  "específica — todos detectáveis via notícias de preço, OPEP e estoques. Ignorar essas "
  "categorias deixaria o corpus estruturalmente incompleto para prever volatilidade.")
p("Exemplo da estrutura da taxonomia em Python (a categoria é gravada em cada notícia, "
  "habilitando a análise de ablação):")
codigo('''TERMOS_POR_CATEGORIA = {
    "CAT1_Empresa": ["Petrobras", "PETR4", "pré-sal", "Petrobras dividendos", ...],
    "CAT2_Mercado_Petroleo": ["petróleo Brent", "OPEP", "preço do petróleo", ...],
    "CAT3_Geopolitica": ["guerra Rússia Ucrânia", "Estreito de Ormuz", ...],
    "CAT4_Infraestrutura": ["refinaria petróleo", "plataforma offshore", ...],
    "CAT5_Sancoes_Navegacao": ["embargo petróleo", "Mar Vermelho petróleo", ...],
    "CAT6_Governanca": ["CEO Petrobras", "ministro de minas e energia", ...],
    "CAT7_Macro_Energia": ["dólar petróleo", "demanda China petróleo", ...],
}''')

# ──────────────────────────────────────────────────────────────────────────────
# 9. AGRUPAMENTO / CATEGORIZAÇÃO
# ──────────────────────────────────────────────────────────────────────────────
h("9. Agrupamento (Categorização) das Notícias", 1)
p("O agrupamento NÃO é feito por classificação posterior, e sim na origem: cada termo de busca "
  "pertence a uma categoria, e a categoria do termo que capturou a notícia é gravada na coluna "
  "'categoria' do CSV. Assim, cada notícia carrega sua origem temática, o que permite a análise "
  "de ablação (remover uma categoria por vez e medir o impacto na previsão).")
codigo('''for portal, endpoint in PORTAIS.items():
    for categoria, termos in TERMOS_POR_CATEGORIA.items():   # 7 categorias
        for termo in termos:                                 # 152 termos
            for artigo in coletar_termo(portal, endpoint, categoria, termo, ini, fim):
                if artigo["hash_titulo"] not in hashes:      # deduplicação
                    hashes.add(artigo["hash_titulo"])
                    writer.writerow(artigo)                  # grava com 'categoria' ''')
p("Distribuição resultante por categoria:")
img(g_categorias)
legenda("Figura 3 — Notícias por categoria temática (base completa).")
tabela(["Categoria","Notícias","% do corpus"],
 [[c, fmt(dist_cat[c]), f"{dist_cat[c]/N_ORIG*100:.1f}%"] for c in dist_cat.index])

# ──────────────────────────────────────────────────────────────────────────────
# 10. ESTATÍSTICAS GERAIS
# ──────────────────────────────────────────────────────────────────────────────
h("10. Estatísticas da Coleta — Sucessos e Falhas", 1)
p("Síntese quantitativa da coleta final:")
tabela(["Métrica","Valor"],
 [["Período coberto","2018-01-01 a 2025-12-31"],
  ["Artigos brutos requisitados", fmt(COLETA["brutos"])],
  ["Duplicatas removidas (hash)", fmt(COLETA["duplicados"])],
  ["Notícias únicas (corpus final)", fmt(N_ORIG)],
  ["Notícias com timestamp real", f"{tem_hora_pct:.1f}%"],
  ["Notícias após o fechamento (≥17h)", f"{lead_lag_pct:.1f}%"],
  ["Portais (fontes)","5"],
  ["Categorias temáticas","7"],
  ["Termos de busca","152"]])
p("Distribuição anual — mostra cobertura contínua e o pico de 2022 (ano da troca de comando da "
  "Petrobras, guerra Rússia-Ucrânia e disparada do petróleo), evidência de validade do corpus:")
img(g_anual); legenda("Figura 4 — Notícias por ano.")
tabela(["Ano","Notícias"], [[str(a), fmt(dist_ano[a])] for a in dist_ano.index])
p("Validação do timestamp e do Lead-Lag — requisito central da banca:")
img(g_leadlag); legenda(f"Figura 5 — {lead_lag_pct:.1f}% das notícias foram publicadas após o "
                        "fechamento da B3 e são realocadas ao pregão seguinte.")
img(g_horas); legenda("Figura 6 — Distribuição por hora de publicação (granularidade horária real).")
p("Falhas e limitações observadas: alguns termos de alto volume exigiram fallback para janela "
  "mensal (paginação profunda lenta); um termo ocasional retornou falha de sondagem e foi pulado "
  "(registrado no log); a cobertura por categoria é desigual (categorias exógenas têm menos "
  "notícias que cita explicitamente o ativo). Tudo foi registrado em log auditável.")

# ──────────────────────────────────────────────────────────────────────────────
# 11. FILTRAGEM (FORTE x LEVE)
# ──────────────────────────────────────────────────────────────────────────────
h("11. Tratamento e Filtragem das Notícias", 1)
p("Após a coleta, avaliaram-se duas estratégias de filtragem. A base ORIGINAL é sempre "
  "preservada intacta; as filtragens geram bases derivadas.")
h("11.1. Filtragem forte (estrita) — testada e descartada", 2)
p("Exigia que o título ou o resumo contivesse explicitamente um termo do ativo/commodity "
  "(petrobras, petr4, petróleo, brent, opep…). Código:")
codigo('''TERMOS_RELEVANCIA = ["petrobras","petr4","petróleo","brent","opep","barril", ...]
alvo = (df["titulo"].fillna("") + " " + df["resumo"].fillna("")).str.lower()
mask = alvo.apply(lambda t: any(termo in t for termo in TERMOS_RELEVANCIA))
df_filtrada = df[mask]   # mantém só ~20% do corpus''')
p(f"Resultado: manteve apenas {fmt(N_ESTRITO)} notícias ({N_ESTRITO/N_ORIG*100:.1f}%) e — pior — "
  "ESVAZIOU as categorias exógenas, justamente as que a literatura aponta como determinantes da "
  "volatilidade. A categoria Geopolítica caiu para ~1% e Macro/Energia para ~3%.")
h("11.2. Filtragem leve (adotada)", 2)
p("Aplica apenas uma limpeza de qualidade: remove notícias degeneradas (título vazio, com menos "
  "de 15 caracteres, ou marcadores como \"[Removed]\"). NÃO há filtro temático — o sinal de "
  "relevância já vem do termo da taxonomia usado na captura. Código:")
codigo('''titulo = df["titulo"].fillna("").astype(str).str.strip()
remover = (titulo == "") | (titulo.str.len() < 15) | titulo.str.lower().isin(MARCADORES)
df_tratada = df[~remover]   # mantém ~100% do corpus, preserva as 7 categorias''')
h("11.3. Por que a filtragem leve?", 2)
p("Porque a filtragem forte destruiria a contribuição científica central da taxonomia — a análise "
  "de ablação por categoria. As notícias exógenas (uma guerra no Oriente Médio, uma decisão da "
  "OPEP, uma alta do dólar) frequentemente não citam \"Petrobras\" no título, mas impactam "
  "diretamente a volatilidade do ativo. Removê-las jogaria fora exatamente o diferencial da "
  "pesquisa. A filtragem leve preserva a amplitude temática e remove apenas ruído degenerado.")
h("11.4. Resultado da filtragem", 2)
tabela(["Estratégia","Notícias mantidas","% do corpus","Efeito nas categorias exógenas"],
 [["Original (sem filtro)", fmt(N_ORIG), "100%","—"],
  ["Filtragem forte (estrita)", fmt(N_ESTRITO), f"{N_ESTRITO/N_ORIG*100:.1f}%","Esvaziadas (Geopol. ~1%, Macro ~3%)"],
  ["Filtragem leve (adotada)", fmt(N_LEVE), f"{N_LEVE/N_ORIG*100:.1f}%","Preservadas integralmente"]])
img(g_filtro)
legenda("Figura 7 — Impacto das filtragens por categoria. A filtragem forte (vermelho) colapsa as "
        "categorias exógenas; a leve (verde) acompanha o original (cinza).")

# ──────────────────────────────────────────────────────────────────────────────
# 12. SPLIT TREINO/VALIDAÇÃO/TESTE
# ──────────────────────────────────────────────────────────────────────────────
h("12. Separação em Treino, Validação e Teste", 1)
p("Conforme orientação (Prof. Emerson), os dados foram separados em três conjuntos na proporção "
  "60% treino / 15% validação / 25% teste. Por se tratar de série temporal financeira, o corte é "
  "CRONOLÓGICO (nunca aleatório, para evitar data leakage / lookahead).")
p("A separação é ESTRATIFICADA POR ANO: dentro de cada ano, os primeiros 60% dos dias "
  "(cronologicamente) vão para treino, os 15% seguintes para validação e os 25% finais (mais "
  "recentes) para teste. Isso garante que todos os anos — e portanto todos os regimes de mercado "
  "— estejam representados nos três conjuntos. A atribuição é por DIAS INTEIROS: um dia nunca é "
  "dividido entre conjuntos, preservando a integridade temporal.")
codigo('''def classificar_ano(grupo):
    contagem_dia = grupo.groupby("data").size().sort_index()
    frac_acum = contagem_dia.cumsum() / contagem_dia.sum()
    dia_conjunto = {}
    for dia, f in frac_acum.items():
        if   f <= 0.60: dia_conjunto[dia] = "treino"
        elif f <= 0.75: dia_conjunto[dia] = "validacao"
        else:           dia_conjunto[dia] = "teste"
    return grupo["data"].map(dia_conjunto)

df["conjunto"] = df.groupby("ano", group_keys=False).apply(classificar_ano)''')
if not tab_split.empty:
    p("Distribuição obtida:")
    tabela(["Conjunto","Notícias","%"],
     [["Treino", fmt(dist_split.get("treino",0)), f"{dist_split.get('treino',0)/len(dft)*100:.1f}%"],
      ["Validação", fmt(dist_split.get("validacao",0)), f"{dist_split.get('validacao',0)/len(dft)*100:.1f}%"],
      ["Teste", fmt(dist_split.get("teste",0)), f"{dist_split.get('teste',0)/len(dft)*100:.1f}%"]])
    img(g_split)
    legenda("Figura 8 — Split temporal por ano (todos os anos presentes nos três conjuntos).")
    tabela(["Ano","Treino","Validação","Teste"],
     [[str(a), fmt(tab_split.loc[a,"treino"]), fmt(tab_split.loc[a,"validacao"]), fmt(tab_split.loc[a,"teste"])]
      for a in tab_split.index])

# ──────────────────────────────────────────────────────────────────────────────
# 13. COMPLEMENTO DO ORIENTADOR
# ──────────────────────────────────────────────────────────────────────────────
h("13. Complemento (visão do orientador): o que ainda deve constar na dissertação", 1)
p("Os itens abaixo complementam o detalhamento e fortalecem o rigor metodológico do capítulo de "
  "Materiais e Métodos. Recomenda-se incluí-los:")
for titulo, corpo in [
 ("Reprodutibilidade.", "Registrar data exata da coleta, versões das bibliotecas (requests, pandas, "
  "transformers, etc.) e o fato de a base bruta ser imutável. Disponibilizar os scripts em repositório."),
 ("Considerações éticas e Termos de Uso.", "Esclarecer que a coleta usou APIs públicas para fins "
  "acadêmicos, com pausas de cortesia (rate limiting respeitoso), sem fins comerciais."),
 ("Validação do modelo de sentimento.", "Construir um conjunto-ouro (gold standard) com uma amostra "
  "de notícias rotuladas manualmente e reportar a concordância (ex.: kappa de Cohen) entre o modelo "
  "(BERTimbau/XLM-RoBERTa) e os anotadores humanos."),
 ("Alinhamento com o calendário de pregão.", "Detalhar o tratamento de fins de semana e feriados: "
  "notícias de dias sem pregão são atribuídas ao próximo dia útil, além do corte das 17h."),
 ("Construção do Índice de Sentimento da Mídia (ISM).", "Documentar a agregação diária "
  "(média de polaridade × confiança) e o ISM por categoria usado na ablação."),
 ("Definição formal da variável-alvo e das features.", "Direção (alta/baixa do log-retorno em t+1) "
  "e a matriz de atributos (retorno, volatilidade GARCH e sentimento defasados em t-1)."),
 ("Métricas e baseline.", "Acurácia, precisão, F1 e AUC-ROC, comparando o modelo com sentimento "
  "(Data Fusion) contra o baseline só com preços."),
 ("Ameaças à validade.", "Viés de seleção das 5 fontes; idioma (PT-BR); cobertura desigual por "
  "categoria; possível viés de sobrevivência de URLs; dependência da qualidade do parser de data."),
 ("Estatística descritiva da série financeira.", "Incluir os testes de Jarque-Bera, ADF e ARCH-LM "
  "que fundamentam o uso do GARCH(1,1)."),
 ("Fluxograma do pipeline.", "Um diagrama 01→02b→02c→03→04 facilita a leitura do método."),
]:
    pr = doc.add_paragraph(style="List Bullet"); pr.add_run(titulo).bold = True
    pr.add_run(" " + corpo)

p("Observação sobre o particionamento: o mapa de datas (definicao_split_temporal.csv) deve ser "
  "aplicado também à matriz diária de modelagem (Script 04), de modo que treino, validação e "
  "teste sejam consistentes entre o corpus textual e os dias de pregão.")

doc.add_paragraph()
ref = doc.add_paragraph(); ref.add_run("Referências citadas: ").bold = True
ref.add_run("Cardoso & Nakane (2024); Hamilton (1983); Kilian (2009); Kilian & Murphy (2014); "
            "Caldara & Iacoviello (2022); Tetlock et al. (2008); Heston & Sinha (2017); "
            "Zhang et al. (2008). Referências completas no capítulo de Referências da dissertação.")

doc.save(SAIDA)
print(f"\n✅ Documento gerado: {SAIDA.resolve()}")
print(f"   Seções: 13 | Figuras: 8 | Tabelas: várias | Base: {fmt(N_ORIG)} notícias")

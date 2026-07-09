# -*- coding: utf-8 -*-
# Gera o Roteiro de Fala (v2) em .docx, alinhado ao deck Apresentacao..._v2.pptx
# (50 slides). Numeracao AUTOMATICA (casa com a ordem real dos slides).
from docx import Document
from docx.shared import Pt, RGBColor

AZUL = RGBColor(0x0b, 0x53, 0x94)
CINZA = RGBColor(0x44, 0x44, 0x44)
VERM = RGBColor(0xb0, 0x30, 0x30)

doc = Document()
st = doc.styles["Normal"].font
st.name = "Calibri"; st.size = Pt(11)
_n = [0]

def h_title(t, sub=""):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = AZUL
    if sub:
        p2 = doc.add_paragraph(); r2 = p2.add_run(sub); r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = CINZA

def slide(titulo, fala, perguntar=None, dica=None):
    _n[0] += 1
    p = doc.add_paragraph(); p.space_before = Pt(10)
    r = p.add_run(f"Slide {_n[0]} — {titulo}"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = AZUL
    pf = doc.add_paragraph(); rf = pf.add_run("Fala: "); rf.bold = True; pf.add_run(fala)
    if dica:
        pd = doc.add_paragraph(); rd = pd.add_run("Como conduzir: "); rd.bold = True; rd.font.color.rgb = CINZA; pd.add_run(dica)
    if perguntar:
        pq = doc.add_paragraph(); rq = pq.add_run("Se perguntarem: "); rq.bold = True; rq.font.color.rgb = VERM; pq.add_run(perguntar)

h_title("Roteiro de Fala — Seminário de Qualificação (v2)",
        "O Impacto do Sentimento de Notícias na Previsão de Direção e Volatilidade da PETR4 · "
        "Vanderlei Barbosa da Silva · PPGIA/PUCPR · 2026 · (50 slides)")
doc.add_paragraph("Instruções rápidas: fale devagar, uma ideia por vez. “Se perguntarem” = resposta pronta "
    "para a banca. Números em negrito são os que valem citar. Mensagem-âncora: o sentimento antecipa "
    "RISCO (volatilidade) muito mais do que direção.").italic = True

# 1
slide("Capa",
 "Bom dia. Sou Vanderlei Barbosa da Silva, orientado pelo Prof. Júlio Nievola e coorientado pelo "
 "Prof. Emerson Paraiso. Vou apresentar minha pesquisa sobre o impacto do sentimento de notícias "
 "na previsão de direção e de volatilidade da ação PETR4.")
# 2
slide("Estrutura da Apresentação",
 "A ordem: introdução e motivação; perguntas, objetivos e hipóteses; contribuições; fundamentação "
 "teórica; trabalhos relacionados; método; resultados preliminares; e cronograma.")
# 3
slide("Introdução", "Começo situando o problema e por que ele é difícil e relevante.")
# 4
slide("Contextualização — o problema",
 "Prever para onde vai um ativo e qual o seu risco é um problema central em finanças: orienta "
 "investimento, gestão de risco e decisão. O preço reflete o passado, mas também notícias e a "
 "economia. E um detalhe: parte da informação aparece primeiro no TEXTO das notícias, antes de virar "
 "cotação. Só que prever é difícil — há muito ruído e eventos extremos.",
 dica="Aqui você planta a semente: 'a informação está no texto antes de virar preço'.")
# 5
slide("Contextualização — por que a PETR4?",
 "Escolhi a PETR4 por quatro razões: é a ação mais negociada e mais noticiada da B3; é duplamente "
 "sensível — sofre com a empresa (governança, dividendos) E com o petróleo (Brent, geopolítica); é "
 "muito politizada, o que torna as notícias informativas; e meu foco é UM ativo, não um índice.")
# 6
slide("Choque Informacional: Geopolítica e a Volatilidade da PETR4",
 "Um exemplo concreto. Na sexta, 27/02, mercado calmo, a PETR4 caiu 1,53%. No sábado, com a bolsa "
 "fechada, os EUA atacam o Irã — um choque geopolítico. Na segunda, o pregão abre em GAP de +4,58% e "
 "o mês fecha em +28%. A notícia se precifica de uma vez, deslocando preço e, principalmente, a "
 "VOLATILIDADE. É esse tipo de evento que sustenta o meu principal achado.",
 dica="Seu slide de ouro — conte como história. Já planta a mensagem 'sentimento → volatilidade'.")
# 7
slide("Motivação — oportunidade e lacuna",
 "A oportunidade vem de dois avanços: a IA de linguagem (Transformers, como o FinBERT) hoje LÊ o "
 "contexto do texto; e o risco é modelável pelo GARCH. A lacuna é unir esses dois mundos — em "
 "português e para um ativo. Minha aposta é fundir sentimento + risco + preço e medir se isso ajuda.")
# 8
slide("Perguntas de Pesquisa",
 "A pergunta principal: como o sentimento extraído por Transformers impacta a previsão de direção e "
 "volatilidade da PETR4? E três secundárias: o texto prevê a direção? o sentimento antecipa o risco? "
 "o mercado reage de forma assimétrica — o viés de negatividade?")
# 9
slide("Objetivos e Hipóteses",
 "O objetivo geral é investigar esse impacto com uma arquitetura de fusão de dados. Os específicos "
 "são as etapas: coletar o corpus, aplicar a taxonomia, extrair o sentimento com FinBERT, modelar o "
 "risco com GARCH e treinar os classificadores com validação cronológica.")
# 10
slide("Hipóteses",
 "Três hipóteses. H1: o sentimento melhora a direção frente a só preços. H2: o sentimento ANTECIPA a "
 "volatilidade. H3: notícias negativas causam reações mais intensas. Adianto: os dados apoiam "
 "fortemente a H2, e a H1 é modesta — vou ser transparente sobre isso.")
# 11
slide("Contribuições",
 "A contribuição principal é um pipeline diário que funde NLP (Transformers) com risco (GARCH) para "
 "a PETR4. As complementares: um corpus datado, evidência do viés de negatividade na B3, e um "
 "framework que projeta o efeito da notícia no pregão seguinte.")
# 12
slide("Contribuições — Framework preditivo",
 "Este esquema mostra o framework: a notícia divulgada após o fechamento é avaliada e o sistema "
 "projeta se ela tende a influenciar a direção do dia seguinte.")
# 13
slide("Fundamentação Teórica", "Agora os conceitos que sustentam o método.")
# 14
slide("Conceitos Fundamentais",
 "Seis conceitos: FinBERT-PT-BR (a IA que lê a manchete e dá um sentimento de −1 a +1); GARCH "
 "(estima o risco e capta que dias turbulentos vêm juntos); XGBoost (meu classificador principal da "
 "direção); SVM (comparação); Data Fusion (juntar sentimento, risco, retorno e categorias num vetor); "
 "e walk-forward (validar no tempo, sem trapaça).",
 perguntar="Por que XGBoost e não deep learning? A base é tabular e de porte médio; boosting é estado da "
 "arte para dados tabulares, treina rápido e mede a importância de cada atributo.")
# 15
slide("Trabalhos Relacionados", "Como cheguei aos trabalhos e o que os diferencia do meu.")
# 16
slide("Protocolo da Revisão Sistemática (RSL)",
 "Fiz uma revisão sistemática em cinco bases (ACM, IEEE, Scopus, Web of Science, CAPES), com termos "
 "em português e inglês. Distingo duas janelas: a LITERATURA, de 2007 (Tetlock) a 2026; e os DADOS, "
 "de 2016 a 2026. O funil PRISMA: 452 identificadas, 423 excluídas, 29 selecionadas.")
# 17
slide("Trabalhos Relacionados — tabela comparativa",
 "Esta é a tabela que compara os principais estudos com o meu. Vou percorrê-la pelas colunas, que "
 "são os critérios que mostram a minha inovação.",
 dica="NÃO leia a tabela inteira. Use o próximo slide para explicar coluna a coluna e volte a esta para apontar.")
# 18
slide("Como ler a tabela comparativa",
 "Coluna por coluna. ATIVO × ÍNDICE: a maioria prevê índices; eu foco um ativo, a PETR4. TEXTO: "
 "trabalhos antigos contam palavras (TF-IDF); eu uso Transformer (FinBERT), que entende contexto e "
 "ironia. IDIOMA: quase tudo é em inglês, só 5 em português; o meu é PT-BR. RISCO: poucos modelam a "
 "volatilidade junto do texto; eu coloco o GARCH na fusão. ALVO: muitos preveem só direção; eu prevejo "
 "direção E volatilidade. A soma dessas colunas — a FUSÃO — é o que nenhum trabalho fez para a PETR4.",
 dica="O slide mais importante da seção. Fale com calma; é aqui que você demonstra a inovação.",
 perguntar="Qual a diferença para o trabalho X? Aponte a coluna: normalmente 'ele usa índice e léxico; eu "
 "uso ativo específico e Transformer, e ainda somo o risco GARCH'.")
# 19
slide("Panorama Quantitativo da Literatura",
 "Em números: 29 estudos de 452, com forte concentração recente. E uma lacuna de idioma: só 5 em "
 "português. Isso reforça o meu diferencial.")
# 20
slide("Lacunas da Literatura — Granularidade",
 "Primeira lacuna, granularidade: sair dos índices inteiros e mirar o risco específico de um ativo "
 "muito politizado, a PETR4.")
# 21
slide("Lacunas da Literatura — Semântica",
 "Segunda, semântica: trocar dicionários estáticos e contagem de palavras por Transformers de ponta, "
 "que entendem contexto, ironia e o jargão do português.")
# 22
slide("Lacunas da Literatura — Arquitetura (Data Fusion)",
 "Terceira, arquitetura: unificar, numa só solução, a IA textual e a predição não linear com o risco "
 "econométrico (GARCH) — o que a literatura ainda trata de forma fragmentada. Estas três lacunas são "
 "o que a minha pesquisa preenche.")
# 23
slide("Método da Pesquisa", "Agora, como faço na prática.")
# 24
slide("Metodologia da Pesquisa",
 "É experimental e aplicada. Experimental porque comparo modelos COM e SEM sentimento, sob as mesmas "
 "condições. Aplicada porque entrego um artefato que gera previsões verificáveis. A unidade é o "
 "pregão: uma previsão por dia. Pressuposto central: causalidade temporal — só uso informação de ANTES.")
# 25
slide("Dados: Preços da PETR4 (B3)",
 "Os preços: cotação diária da PETR4, 2016 a 2026, 2.612 pregões, via B3/yfinance. Derivo o "
 "log-retorno e defino o alvo: retorno positivo é 'alta' (1), caso contrário 'baixa' (0).",
 perguntar="Por que log-retorno? Estabiliza a série (estacionariedade) e torna as variações comparáveis no tempo.")
# 26
slide("Dados: Corpus de Notícias",
 "O corpus: 261.960 manchetes datadas, 2016 a 2026, de cinco portais, via API WordPress, com data e "
 "HORA exatas. Cada notícia recebe uma categoria. As do dia formam o índice de sentimento; as de após "
 "as 17h contam para o pregão seguinte (Lead-Lag).")
# 27 (NOVO)
slide("Coleta com data e hora (WordPress REST API)",
 "Um ponto que a banca valoriza: como garanto a data e a HORA de cada notícia. Em vez de APIs "
 "comerciais, de janela curta, uso a API REST do WordPress dos cinco portais — o endereço "
 "/wp-json/wp/v2/posts. Ela devolve, para cada notícia, o horário de Brasília e o UTC. Isso dá "
 "marcação temporal precisa, que é o que viabiliza o Lead-Lag: separar o que saiu antes e depois do "
 "fechamento das 17h. E, com cinco fontes, não dependo de um único portal.",
 dica="Se a banca já cobrou o timestamp antes, ESTE slide é a sua resposta direta.")
# 28
slide("Arquitetura da Solução",
 "Este é o desenho geral. Vou explicá-lo por etapas no próximo slide, caixa por caixa.",
 dica="Não explique o diagrama inteiro aqui; use o slide seguinte, com as 5 etapas.")
# 29
slide("Arquitetura em 5 etapas",
 "Cinco etapas. Um: COLETA de notícias e preços com data e hora. Dois: SENTIMENTO — o FinBERT lê cada "
 "manchete e o dia vira o ISM. Três: RISCO — o GARCH estima a volatilidade. Quatro: FUSÃO — junto "
 "sentimento, risco, retorno e categorias num vetor, tudo em t−1. Cinco: CLASSIFICAÇÃO — o XGBoost "
 "projeta a direção do próximo pregão, com walk-forward.",
 dica="Conte nos dedos: 1 coleta, 2 sentimento, 3 risco, 4 fusão, 5 classificação.")
# 30
slide("Fusão de Dados e Predição da Direção",
 "Data Fusion é juntar informações de naturezas diferentes num único vetor por pregão: sentimento, "
 "risco (GARCH), retorno e categorias — tudo do dia anterior (t−1). Esse vetor entra no classificador, "
 "que devolve UMA previsão de direção. A pergunta: esse vetor fundido prevê melhor que só preços?",
 perguntar="Por que t−1? Para não usar o futuro. Só entra o que já era conhecido antes do pregão previsto.")
# 31
slide("Taxonomia Temática das Notícias (7 categorias)",
 "Em vez de deixar um algoritmo adivinhar os temas, defini a priori 7 categorias com 152 termos, com "
 "base na literatura econômica. Cada categoria tem uma via de impacto no preço. É interpretável e "
 "auditável — qualquer um confere.")
# 32 (NOVO)
slide("Análise de sentimento — do texto ao número",
 "Como transformo texto em número. Uso o FinBERT-PT-BR, um BERT ajustado para finanças em português "
 "— troquei um classificador genérico por um especializado. Para cada manchete, ele dá um rótulo "
 "(positivo, negativo ou neutro) e uma confiança, que viram um índice de −1 a +1. Exemplo: 'Petrobras "
 "anuncia dividendos recordes' é positivo; 'ataque a porto petroleiro reduz a oferta' é negativo para "
 "o mercado, mas tende à alta para a produtora. Os 152 termos — Brent, OPEP, Estreito de Ormuz, "
 "dividendos — guiam a relevância. No fim do dia, agrego tudo no ISM.",
 perguntar="O FinBERT foi treinado nos seus dados? Não — ele já vem pré-treinado em finanças em português; "
 "eu apenas o APLICO às manchetes. Isso evita sobreajuste e mantém a reprodutibilidade.")
# 33
slide("Notícias Divergentes no Mesmo Pregão",
 "E quando há notícia boa e ruim no mesmo dia? Cada polaridade é ponderada pela confiança do modelo, "
 "e o saldo forma o ISM. Assim, mesmo com sinais opostos, resulta UM número por pregão.")
# 34 (NOVO)
slide("Volatilidade: o modelo GARCH(1,1)",
 "Aqui explico o risco. A volatilidade — o tamanho das oscilações — é estimada por um GARCH(1,1). Ele "
 "capta o agrupamento: dias turbulentos vêm em sequência. Uso a versão t-Student porque os retornos "
 "da PETR4 têm caudas pesadas — Jarque-Bera rejeita a normalidade — e há efeito ARCH, confirmado pelo "
 "ARCH-LM com p<0,001. Ou seja, o GARCH é justificado pelos dados, não é arbitrário. O gráfico mostra "
 "a volatilidade ao longo do tempo — os picos são os períodos de estresse.",
 dica="Aponte os picos do gráfico: 'é nesses momentos que o sentimento das notícias mais informa'.")
# 35
slide("Separação dos Dados (validação cronológica)",
 "Como divido os dados: nunca aleatório, sempre no tempo. Três blocos em sequência: TREINO, 60%, "
 "1.566 pregões (2016–2022), onde o modelo aprende; VALIDAÇÃO, 15%, 391 pregões (2022–2023), onde "
 "ajusto hiperparâmetros e limiar; e TESTE, 25%, 653 pregões (2023–2026), consultado UMA vez no fim. "
 "O walk-forward repete isso em janelas deslizantes.",
 perguntar="O que ficou para validação? O bloco do MEIO — 2022 a 2023, 391 pregões (15%). Fica ENTRE treino "
 "e teste e calibra o modelo sem tocar no teste.")
# 36
slide("Algoritmo 1 — construção da base diária",
 "Este algoritmo constrói a base. Leia como receita — aponte os dois laços. 'Para cada notícia': "
 "categoria pela taxonomia, sentimento pelo FinBERT e, se saiu após 17h, vai para o dia seguinte. "
 "'Para cada pregão': resumo o sentimento no ISM (ponderado pela confiança), calculo o log-retorno, "
 "estimo a volatilidade com o GARCH, defino o rótulo e monto o vetor — tudo de t−1. Sai a base D.",
 dica="Não leia linha por linha. Aponte os dois 'para cada'. As setas ▷ são só comentários.")
# 37
slide("Algoritmo 2 — treino e avaliação walk-forward",
 "Este treina e avalia. Divido no tempo (60/15/25). Em janelas walk-forward, treino o XGBoost no "
 "passado e ajusto hiperparâmetros e o limiar na validação. Com o modelo final, para cada pregão do "
 "teste calculo a probabilidade de alta e decido pelo limiar. Comparo com o real UMA vez e com o "
 "baseline de ~53%. No fim: acurácia, F1, AUC e importância dos atributos.",
 perguntar="O que é o limiar δ? É o ponto de corte da probabilidade. Em vez de 0,50 fixo, eu o calibro na "
 "validação (aqui, 0,46) para equilibrar os acertos entre alta e baixa.")
# 38
slide("Métricas de Avaliação",
 "Como avalio. A saída é a direção do dia seguinte. Uso acurácia (só vale se superar o baseline de "
 "~53%), precisão, revocação e F1 (que junta as duas e desmascara viés de classe). E o AUC, que "
 "explico a seguir. Tudo no split cronológico 60/15/25.",
 perguntar="O que é AUC? É a qualidade do RANKING, de 0 a 1: a chance de o modelo dar nota maior a uma alta "
 "real do que a uma baixa real. 0,50 é acaso; o meu fica ~0,51–0,52, coerente com a leitura honesta.")
# 39
slide("Resultados Preliminares", "Agora os resultados — e vou ser transparente com os números.")
# 40
slide("Análise Exploratória – Filtragem de Ruídos",
 "Primeiro achado: das 261.960 notícias, ~85% são ruído — não deslocam o preço de forma anormal. O "
 "sinal se concentra nos ~15% que coincidem com rompimentos de estresse no GARCH. Poucas notícias, "
 "mas potentes, carregam a informação de risco.")
# 41
slide("Experimentos e Resultados: Direção",
 "Na direção, o melhor é o XGBoost com Data Fusion: 52,22% de acurácia e AUC 0,514. O sentimento traz "
 "+2,45 pontos sobre só preços, e é o atributo mais importante. Mas sou honesto: o ganho é MODESTO e "
 "não significativo (p=0,145), e não supera o baseline de ~53%. A ablação mostra que "
 "'Sanções/Navegação' é a categoria mais informativa.",
 perguntar="O que é ablação? Tiro UMA categoria, treino de novo e vejo quanto a acurácia CAI. Quanto mais "
 "cai, mais aquela categoria importava.")
# 42 (NOVO)
slide("Ablação por categoria temática",
 "Esta tabela mede a importância de cada categoria por ablação. O modelo completo faz 53,45%. A que "
 "mais faz falta é 'Acordos, Sanções e Navegação': sem ela, cai 3,68 pontos. Depois vêm Governança e "
 "Macroeconomia. Faz sentido econômico: sanções e navegação mexem na oferta física do petróleo.",
 perguntar="Por que Sanções é a mais importante? Porque afeta o fluxo físico da commodity (Kilian, 2009) — "
 "e a Petrobras, como produtora, é sensível a choques de oferta.")
# 43
slide("Comparação com o baseline",
 "Coração da honestidade da pesquisa. O baseline — chutar sempre 'alta' — acerta ~53,1%. Meu modelo "
 "faz 52,22%: NÃO supera o baseline, e tudo bem, porque direção diária é difícil (mercados eficientes). "
 "O ganho REAL está na VOLATILIDADE: o Granger mostra o sentimento a antecipando com p<0,001; e no "
 "efeito assimétrico, nos piores dias o sentimento eleva o retorno em +261 bps. Mensagem: o sentimento "
 "é um radar de RISCO, não uma bússola de direção.",
 dica="Se perguntarem 'a acurácia é baixa', ESTE é o slide para responder — com confiança, não como desculpa.",
 perguntar="Por que não bate o baseline? Direção diária é quase um passeio aleatório. O valor não é bater a "
 "direção, e sim mostrar — com significância — que o sentimento antecipa o RISCO.")
# 44
slide("Modelagem Preditiva Direcional",
 "Resumindo os três resultados: direção com ganho modesto e não significativo; efeito assimétrico "
 "forte nos piores dias (viés de negatividade); e sentimento → volatilidade altamente significativo. "
 "É a evidência que sustenta a H2.")
# 45 (NOVO)
slide("Sentimento × Volatilidade",
 "Este gráfico relaciona o sentimento diário com a volatilidade. É a evidência visual por trás do meu "
 "principal achado: quando o sentimento se intensifica, a volatilidade responde. É o que o Granger "
 "confirma — sentimento antecipa volatilidade com p<0,001. Reforça a mensagem: o sentimento é um radar "
 "de risco.",
 dica="Conecte ao slide da comparação com o baseline — é a mesma mensagem, agora visual.")
# 46
slide("Discussão dos Resultados",
 "O grande desafio é a baixa relação sinal-ruído — 85% de ruído. O sinal está nos 15% de choque. Na "
 "direção, o ganho é modesto (mercados eficientes). Na volatilidade e na assimetria, os resultados são "
 "fortes e significativos. E as séries são robustas: caudas pesadas, estacionárias e com efeito ARCH — "
 "o GARCH é justificado. Conclusão: para a PETR4, o sentimento é mais radar de risco do que bússola.")
# 47
slide("Cronograma de Execução", "Por fim, o cronograma até a defesa.")
# 48
slide("Próximas etapas",
 "Refinar a coleta com timestamps, reprocessar sentimento e GARCH na janela ampliada, fazer o tuning "
 "dos classificadores, consolidar resultados, redigir o documento final e submeter um artigo. A defesa "
 "está prevista para março de 2027.")
# 49
slide("Obrigado", "Obrigado pela atenção. Fico à disposição para perguntas, sugestões e comentários.")
# 50
slide("Fim", "(Slide de encerramento.)")

doc.add_page_break()
h_title("Perguntas difíceis — respostas prontas")
dif = [
 ("A acurácia de 52% não é fraca?",
  "Para direção diária de ação, 52% já fica no teto do que a literatura alcança — o mercado é quase "
  "eficiente. E eu não escondo: a contribuição forte e significativa é na volatilidade (Granger p<0,001) "
  "e no viés de negatividade (quantílica, +261 bps nos piores dias)."),
 ("Isso serviria para ganhar dinheiro? (backtest)",
  "Fiz um backtest honesto com custo de 10 bps por troca: a estratégia rendeu 19,3% contra 44% do "
  "buy-and-hold. Ou seja, NÃO recomendo como trading — o valor é científico e de gestão de risco."),
 ("Por que FinBERT e não um LLM/ChatGPT?",
  "O FinBERT-PT-BR é especializado em finanças e português, é reprodutível, roda localmente e é "
  "auditável. Um LLM genérico seria caro, não determinístico e difícil de reproduzir."),
 ("Como garante que não há vazamento de dados?",
  "Split cronológico (nunca aleatório), atributos sempre em t−1, scaler ajustado só no treino e teste "
  "consultado uma única vez. O walk-forward reforça isso."),
 ("O sentimento realmente ajuda ou é ruído?",
  "Na direção, o ganho é modesto (+2,45 pp, não significativo). Mas na volatilidade o Granger é "
  "significativo em TODAS as defasagens (p<0,001) — evidência robusta, não impressão."),
 ("Por que 'Sanções/Navegação' é a categoria mais informativa?",
  "Porque afeta o fluxo físico do petróleo (Kilian, 2009). Como a Petrobras é produtora, choques de "
  "oferta a atingem diretamente — por isso essa categoria pesa mais na ablação (−3,68 pp)."),
]
for q, a in dif:
    p = doc.add_paragraph(); r = p.add_run("• " + q); r.bold = True
    doc.add_paragraph(a)

doc.save("Roteiro_Falas_v2.docx")
print("Salvo: Roteiro_Falas_v2.docx | slides narrados:", _n[0])

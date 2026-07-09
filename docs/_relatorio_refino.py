# -*- coding: utf-8 -*-
# Gera o relatório Word (docx) do refino pedido pela banca (jul/2026),
# lendo os datasets REAIS gerados. Saída: docs/Relatorio_Refino_Banca_Jul2026.docx
import json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RAIZ = Path(__file__).resolve().parents[1]
DS = RAIZ / "datasets_refino"
MP = RAIZ / "Mestrado_PETR4"
AZUL = RGBColor(0x0b, 0x53, 0x94); CINZA = RGBColor(0x55, 0x55, 0x55); VERD = RGBColor(0x1a, 0x7f, 0x37)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)

def H(t, size=16, cor=AZUL, before=12):
    p = doc.add_paragraph(); p.space_before = Pt(before)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = cor; return p
def P(t, bold=False, italic=False, cor=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if cor: r.font.color.rgb = cor
    return p
def B(t):
    doc.add_paragraph(t, style="List Bullet")
def tabela(cab, linhas, larg=None):
    t = doc.add_table(rows=1, cols=len(cab)); t.style = "Light Grid Accent 1"
    for j, c in enumerate(cab):
        cell = t.rows[0].cells[j]; cell.text = ""
        r = cell.paragraphs[0].add_run(c); r.bold = True; r.font.size = Pt(9.5)
    for ln in linhas:
        cells = t.add_row().cells
        for j, v in enumerate(ln):
            cells[j].text = ""; rr = cells[j].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9.5)
    if larg:
        for j, w in enumerate(larg):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    return t

# ── Capa ──────────────────────────────────────────────────────────────────────
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Relatório de Refino — Ponderações da Banca"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZUL
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Seminário de julho/2026 · Dissertação PETR4 · Vanderlei Barbosa da Silva · PPGIA/PUCPR")
r.italic = True; r.font.color.rgb = CINZA
P("")
P("Princípio: nada foi inventado — cada dado cita a fonte (arquivo real da pesquisa ou URL, "
  "no caso do item 8, o único autorizado a usar fontes externas). Este documento resume tudo "
  "o que foi feito hoje, para leitura e avaliação.", italic=True)

# ── Sumário de itens ─────────────────────────────────────────────────────────
H("1. Sumário das ponderações e status")
tabela(["Item", "Ponderação", "O que foi feito", "Status"], [
 ["1", "Por que não Bloomberg Línea?", "Diagnóstico técnico (Arc Publishing, sem WordPress)", "Respondido"],
 ["2", "Resultados sem volatilidade", "Dataset de resultados de volatilidade consolidado", "Feito"],
 ["3", "Direção pior que a moeda", "Leitura honesta + caminho de melhoria", "Respondido"],
 ["4", "Refinar dataset de notícias", "01 (após-17h) e 02 (enriquecido) versionados", "Feito"],
 ["5", "Dataset da RSL", "25 estudos, 14 colunas (com fonte das notícias)", "Feito"],
 ["7", "Pasta de datasets", "Criada datasets_refino/ (evita colisão de nome)", "Feito"],
 ["8", "Encoders melhores?", "Pesquisa + experimento de fine-tuning pronto", "Feito"],
], larg=[0.5, 2.2, 3.2, 0.9])

# ── Item 4 · datasets ────────────────────────────────────────────────────────
H("2. Item 4 — Refino do dataset de notícias (datasets versionados)")
P("Gerados por datasets_refino/gerar_datasets_refino.py (reprodutível), só de dados reais:")
try:
    n_apos = len(pd.read_csv(DS / "01_noticias_apos_17h_v1.csv", usecols=[0]))
except Exception:
    n_apos = "—"
tabela(["Arquivo", "Conteúdo", "Linhas"], [
 ["01_noticias_apos_17h_v1.csv", "Notícias publicadas após 17h (Lead-Lag)", f"{n_apos:,}" if isinstance(n_apos,int) else n_apos],
 ["02_..._enriquecido_v1.csv", "Notícia + sentimento (FinBERT) + volatilidade (GARCH) + parâmetros dos encoders", f"{n_apos:,}" if isinstance(n_apos,int) else n_apos],
], larg=[2.6, 3.6, 0.8])
B("4A — filtro pela hora de Brasília (coluna Data_Coleta ≥ 17h). São 26,4% do corpus.")
B("4B — volatilidade casada ao PRÓXIMO pregão real (merge_asof); 0 lacunas.")
B("4C — versionamento _vN: cada refinamento gera novo arquivo, sem apagar os anteriores (comparação).")

# ── Item 2 · volatilidade ────────────────────────────────────────────────────
H("3. Item 2 — Resultados de VOLATILIDADE (o ponto forte)")
P("A banca observou que os resultados enfatizaram a direção. A volatilidade já existia e foi "
  "consolidada em datasets_refino/03_resultados_volatilidade_v1.csv:")
try:
    v = pd.read_csv(DS / "03_resultados_volatilidade_v1.csv")
    dest = v[v["Analise"].str.contains("VOLATILIDADE|quantílica|regime", case=False, na=False)].head(9)
    tabela(list(dest.columns[:5]), dest.iloc[:, :5].values.tolist(), larg=[2.1, 1.6, 1.0, 0.8, 0.8])
except Exception as e:
    P(f"(não foi possível ler o CSV: {e})")
B("Granger sentimento→volatilidade: significativo em todas as defasagens (p ≤ 0,0002).")
B("Quantílica: τ=0,05 → +261 bps (p=0,034); τ=0,25 → +121 bps (p=0,025) — viés de negatividade.")

# ── Item 3 · direção ─────────────────────────────────────────────────────────
H("4. Item 3 — A direção ficou 'pior que a moeda' (leitura honesta)")
try:
    m = pd.read_csv(MP / "resultados_modelos_petr4.csv")
    linhas = [[r["Modelo"], f'{r["Acurácia"]}%', f'{float(r["AUC-ROC"]):.3f}'] for _, r in m.iterrows()]
    tabela(["Modelo", "Acurácia", "AUC"], linhas, larg=[3.6, 1.2, 1.0])
except Exception:
    pass
B("49,77% (XGBoost só-preços) é estatisticamente indistinguível de 50% — é ≈ acaso, não 'pior'.")
B("Isso é esperado: direção diária é quase passeio aleatório (mercados eficientes); 52–56% é o teto da literatura.")
B("O sentimento agrega +2,45 pp, mas o ganho direcional não é significativo (p=0,145). O ganho real é na volatilidade.")

# ── Item 5 · RSL ─────────────────────────────────────────────────────────────
H("5. Item 5 — Dataset da Revisão Sistemática (25 estudos)")
P("datasets_refino/04_revisao_sistematica_estudos_v1.csv — 14 colunas, incluindo, como pedido, "
  "ONDE cada estudo capturou as notícias e COMO. Exemplos (extraídos dos PDFs):")
try:
    rsl = pd.read_csv(DS / "04_revisao_sistematica_estudos_v1.csv")
    amostra = rsl[rsl["#"].isin([1, 2, 4, 16, 24, 25])]
    linhas = [[f'{r["Autor(es)"]} ({r["Ano"]})', r["Idioma"], str(r["Fonte_Noticias"])[:52]] for _, r in amostra.iterrows()]
    tabela(["Estudo", "Idioma", "Fonte das notícias / como coletaram"], linhas, larg=[1.9, 0.9, 3.6])
except Exception as e:
    P(f"(erro ao ler RSL: {e})")
B("Preenchidos com fonte: autor, ano, título, idioma, veículo, objetivo, método, encoder, fonte das notícias, método de coleta.")
B("Resultados: preenchidos quando constam no resumo (ex.: Bollen 87,6%; Narde 95,1%); parâmetros detalhados → 'ver artigo'.")
P("⚠️ Inconsistência a resolver: o Cap. 2 diz 25 estudos; o slide do seminário dizia 29. Uniformizar.", cor=RGBColor(0xb0,0x30,0x30))

# ── Item 1 · Bloomberg ───────────────────────────────────────────────────────
H("6. Item 1 — Bloomberg Línea como fonte?")
B("Não usa WordPress: /wp-json → HTTP 404. O robots.txt revela Arc Publishing e bloqueia a API /pf/api/v3/*.")
B("Só há um Google-News sitemap com ~50 artigos das últimas 48h → inviável para o corpus histórico (2016–2026).")
B("Encaminhamento: registrar como trabalho futuro (coletor forward via sitemap Arc), respeitando robots.txt e o paywall/termos.")
P("Fonte: https://www.bloomberglinea.com.br/robots.txt", italic=True, cor=CINZA)

# ── Item 8 · encoders ────────────────────────────────────────────────────────
H("7. Item 8 — Encoders melhores (pesquisa externa) + experimento")
tabela(["Opção", "O que é", "Observação"], [
 ["Albertina PT-BR (900M)", "Encoder DeBERTa, SOTA para PT", "Mais forte que BERTimbau; não é de finanças → fine-tuning"],
 ["BERTimbau-large (330M)", "Versão maior do modelo-base", "Upgrade barato para testar"],
 ["turing-usp/FinBertPTBR", "FinBERT-PT alternativo", "Comparação direta com o atual"],
 ["LLMs (FinGPT/GPT-4o)", "Zero/few-shot", "Não determinístico/caro → só baseline"],
], larg=[1.8, 2.2, 2.4])
P("Recomendação: fine-tunar o Albertina PT-BR e comparar ao FinBERT-PT-BR sob o mesmo protocolo. "
  "Ressalva honesta: um encoder melhor tende a ajudar mais a VOLATILIDADE do que a direção "
  "(o gargalo da direção é a eficiência do mercado).")
P("Experimento pronto: src/sentimento/12_finetune_albertina_ptbr.py — usa o conjunto-ouro rotulado "
  "por humanos e compara acurácia, F1-macro e Kappa de Cohen com o FinBERT, no mesmo teste. "
  "Requer GPU + 'pip install datasets'.", bold=False)
P("Fontes: huggingface.co/PORTULAN/albertina-900m-portuguese-ptbr-encoder · arXiv 2305.06721 · "
  "huggingface.co/lucas-leme/FinBERT-PT-BR · huggingface.co/turing-usp/FinBertPTBR", italic=True, cor=CINZA)

# ── Rodada de experimentos (execução autônoma) ──────────────────────────────
H("8. Rodada de experimentos de data fusion (execução autônoma, 04/07/2026)")
P("Suíte reprodutível (src/modelagem/13_experimentos_datafusion.py) replicando as técnicas de "
  "maior desempenho da literatura: stacking (Barak 2017), tuning/limiar (Nobre 2019), sentimento "
  "por categoria (Nguyen 2015), RF/AUC (Ballings 2015), walk-forward (Oliveira 2017). Split "
  "cronológico 60/15/25 (653 pregões de teste). A cada teste, um dataset foi salvo (data+modelo+atributos).")
try:
    ex = pd.read_csv(DS / "resultados_experimentos_datafusion_2026-07-04.csv").sort_values("acuracia_teste", ascending=False).head(6)
    linhas = [[r["modelo"], r["atributos"], f'{r["acuracia_teste"]}%', f'{r["auc"]:.3f}',
               "sim" if r["supera_baseline"] else "não"] for _, r in ex.iterrows()]
    tabela(["Modelo", "Atributos", "Acurácia (teste)", "AUC", "> baseline"], linhas, larg=[1.6, 1.4, 1.5, 1.0, 1.0])
except Exception as e:
    P(f"(erro ao ler experimentos: {e})")
P("Melhor: XGBoost (3 atributos-base) = 54,52% — supera o baseline (53,14%) e é significativo vs. "
  "acaso (binomial p=0,012). Confirma o 54,5% da dissertação. Stacking e conjunto completo NÃO "
  "superaram (sobreajuste); ganho direcional modesto — o forte segue sendo a volatilidade.", cor=VERD)
P("Fine-tuning de ENCODER: BLOQUEADO honestamente — sem GPU (CUDA indisponível) e o conjunto-ouro "
  "está com 0/300 rótulos humanos. O script (src/sentimento/12_finetune_albertina_ptbr.py) está "
  "pronto e sai com aviso claro até a rotulagem ser feita. Pendência: rotular ≥200 manchetes + GPU.",
  cor=RGBColor(0xb0, 0x30, 0x30))

# ── Filtro de relevância (execução autônoma) ────────────────────────────────
H("9. Filtro de relevância — pipeline reordenado (05/07/2026)")
P("Testou-se a hipótese: coleta → após-17h → FILTRO DE RELEVÂNCIA → predição. Descobriu-se que o "
  "pipeline atual NÃO filtra relevância (o ISM diário é a média de TODAS as notícias). Reconstruiu-se "
  "o sinal com 3 critérios e re-treinou-se o XGBoost (2018–2025, split 60/15/25).")
try:
    rr = pd.read_csv(DS / "resultados_relevancia_2026-07-05.csv")
    nm = {"todas_apos17h": "Todas (após 17h)", "relevante_direta": "Relevantes (menção direta)", "categoria_CAT1_CAT2": "CAT1+CAT2"}
    linhas = [[nm.get(r["variante"], r["variante"]), f'{r["cobertura_dias_com_noticia"]:.2f}',
               f'{r["acuracia_teste"]}%', f'{r["auc"]:.3f}', f'{r["p_granger_sent_para_vol_min"]:.4f}'] for _, r in rr.iterrows()]
    tabela(["Sinal de sentimento", "Cobertura", "Acurácia", "AUC", "p Granger (→vol)"], linhas, larg=[2.4, 1.1, 1.1, 1.0, 1.3])
except Exception as e:
    P(f"(erro: {e})")
P("Achado honesto: o filtro de relevância NÃO elevou a acurácia direcional (52,11% com e sem filtro; "
  "AUC subiu de leve 0,521→0,525; CAT1+CAT2 até piorou). Motivo: cobertura ~100% (todo dia tem notícia "
  "relevante) + corpus já filtrado pela taxonomia + gargalo é a eficiência do mercado, não o ruído. "
  "A relação sentimento→volatilidade segue significativa (Granger p<0,02) em todas as variantes.", cor=VERD)
P("Sugestões de orientador p/ acurácia: (i) alvo = VOLATILIDADE (magnitude); (ii) rótulo com zona morta "
  "(descartar dias ~0); (iii) modelo em 2 estágios (detectar choque informacional, depois direção); "
  "(iv) sentimento direcionado ao ativo (aspect-based, Kilian/Hamilton); (v) ponderar por reação de "
  "mercado; (vi) encoder Albertina (pende rótulos+GPU); (vii) intradiário.")
P("CSVs gerados no seu esquema exato: 10_noticias_apos_17h, 11_noticias_relevantes, "
  "12_relevantes_{Treino,Validacao,Teste} (com sentimento, risco, predição, encoder/params e data).", italic=True, cor=CINZA)

# ── Próximos passos ──────────────────────────────────────────────────────────
H("10. Próximos passos sugeridos")
B("Definir o número oficial de estudos da RSL (25 vs 29) e uniformizar deck + dissertação.")
B("Rodar o experimento de fine-tuning (Albertina) em GPU e comparar ao FinBERT-PT-BR.")
B("Promover os resultados de volatilidade ao corpo dos Resultados (tabela/figura dedicada).")
B("Completar Parâmetros/Resultados por estudo na RSL via leitura profunda dos PDFs (sob demanda).")

P("")
P("Detalhamento completo em docs/RESPOSTAS_BANCA_JUL2026.md. Datasets em datasets_refino/.", italic=True, cor=CINZA)

doc.save(str(RAIZ / "docs" / "Relatorio_Refino_Banca_Jul2026.docx"))
print("✓ docs/Relatorio_Refino_Banca_Jul2026.docx gerado.")

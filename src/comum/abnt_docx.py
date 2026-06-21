# -*- coding: utf-8 -*-
# ==============================================================================
#   abnt_docx.py — Módulo reutilizável de formatação ABNT para documentos Word
#   Dissertação PETR4 | Vanderlei Barbosa da Silva
#
#   Centraliza a formatação ABNT (NBR 14724, 6023, 6024, 10520) para que TODOS
#   os documentos de etapa (coleta, sentimento, modelagem) tenham padrão idêntico.
#
#   Normas aplicadas:
#     • Fonte Times New Roman 12 (corpo), 10 (citações longas, fontes, código)
#     • Espaçamento 1,5 no corpo; simples em citações/fontes/quadros
#     • Margens: esquerda/superior 3 cm, direita/inferior 2 cm
#     • Texto justificado, recuo de primeira linha 1,25 cm
#     • Seções numéricas progressivas (NBR 6024): primária em CAIXA ALTA negrito
#     • Tabelas abertas (sem bordas verticais), título acima e Fonte abaixo
#     • Quadros (código/texto) fechados, título acima e Fonte abaixo
#     • Figuras: título acima, Fonte abaixo
#     • Citações autor-data (NBR 10520) e Referências (NBR 6023)
# ==============================================================================

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONTE       = "Times New Roman"
FONTE_MONO  = "Consolas"
ANO_FONTE   = "Elaborado pelo autor (2026)"


# ─── Documento base ───────────────────────────────────────────────────────────
def novo_documento():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(3); sec.left_margin = Cm(3)
        sec.bottom_margin = Cm(2); sec.right_margin = Cm(2)
    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    return doc


# ─── Capa / folha de rosto simplificada ───────────────────────────────────────
def capa(doc, titulo, subtitulo, autor, orientador, instituicao, descricao=None):
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(titulo.upper()); r.bold = True; r.font.size = Pt(14); r.font.name = FONTE
    if subtitulo:
        s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = s.add_run(subtitulo); rs.italic = True; rs.font.size = Pt(12); rs.font.name = FONTE
    for _ in range(2):
        doc.add_paragraph()
    for txt in (autor, orientador, instituicao):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(txt); rr.font.size = Pt(12); rr.font.name = FONTE
    if descricao:
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(8)
        rr = p.add_run(descricao); rr.font.size = Pt(10); rr.font.name = FONTE
    doc.add_page_break()


# ─── Seções numeradas (NBR 6024) ──────────────────────────────────────────────
def secao(doc, numero, texto, nivel=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if nivel == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    rotulo = f"{numero} {texto.upper()}" if nivel == 1 else f"{numero} {texto}"
    r = p.add_run(rotulo)
    r.bold = True; r.font.name = FONTE; r.font.size = Pt(12)
    return p


# ─── Parágrafo de corpo (justificado, recuo 1,25 cm) ──────────────────────────
def paragrafo(doc, texto, recuo=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if recuo:
        pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    # Permite trechos em negrito via marcador **...**
    partes = texto.split("**")
    for i, parte in enumerate(partes):
        run = p.add_run(parte)
        run.font.name = FONTE; run.font.size = Pt(12)
        if i % 2 == 1:
            run.bold = True
    return p


# ─── Citação longa (NBR 10520: recuo 4 cm, fonte 10, espaço simples) ──────────
def citacao_longa(doc, texto, fonte_citacao):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    r = p.add_run(texto); r.font.name = FONTE; r.font.size = Pt(10)
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rc = pc.add_run(fonte_citacao); rc.font.name = FONTE; rc.font.size = Pt(10)
    return p


def _legenda(doc, texto, acima=True, fonte=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6 if acima else 2)
    p.paragraph_format.space_after = Pt(2 if acima else 6)
    r = p.add_run(texto)
    r.font.name = FONTE; r.font.size = Pt(10 if not fonte else 10)
    if fonte:
        r.italic = False
    return p


# ─── Bordas ───────────────────────────────────────────────────────────────────
def _bordas_tabela_aberta(table):
    """ABNT: tabela aberta — bordas só no topo, base e sob o cabeçalho."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "8"); el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none"); el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)
    # Linha sob o cabeçalho
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8"); bot.set(qn("w:color"), "000000")
        tcB.append(bot); tcPr.append(tcB)


def _fmt_celula(cell, negrito=False, tamanho=11):
    for par in cell.paragraphs:
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        par.paragraph_format.space_after = Pt(0)
        for r in par.runs:
            r.font.name = FONTE; r.font.size = Pt(tamanho); r.bold = negrito


def tabela_abnt(doc, numero, titulo, cabecalho, linhas, fonte=None):
    """Tabela ABNT: 'Tabela N – título' acima, tabela aberta, 'Fonte:' abaixo."""
    _legenda(doc, f"Tabela {numero} – {titulo}")
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, c in enumerate(cabecalho):
        t.rows[0].cells[i].text = str(c)
        _fmt_celula(t.rows[0].cells[i], negrito=True)
    for lin in linhas:
        cells = t.add_row().cells
        for i, v in enumerate(lin):
            cells[i].text = str(v); _fmt_celula(cells[i])
    _bordas_tabela_aberta(t)
    _legenda(doc, f"Fonte: {fonte or ANO_FONTE}.", acima=False, fonte=True)
    doc.add_paragraph()
    return t


def quadro_codigo(doc, numero, titulo, codigo, fonte=None):
    """Quadro ABNT (fechado) para listagem de código: título acima, Fonte abaixo."""
    _legenda(doc, f"Quadro {numero} – {titulo}")
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(codigo.strip("\n"))
    run.font.name = FONTE_MONO; run.font.size = Pt(9)
    cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F4F4F4")
    cell._tc.get_or_add_tcPr().append(shd)
    _legenda(doc, f"Fonte: {fonte or ANO_FONTE}.", acima=False, fonte=True)
    doc.add_paragraph()
    return t


def figura_abnt(doc, numero, titulo, caminho_img, largura_cm=15, fonte=None):
    """Figura ABNT: 'Figura N – título' acima, imagem centralizada, 'Fonte:' abaixo."""
    from docx.shared import Cm as _Cm
    from pathlib import Path
    _legenda(doc, f"Figura {numero} – {titulo}")
    if caminho_img and Path(caminho_img).exists():
        doc.add_picture(str(caminho_img), width=_Cm(largura_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _legenda(doc, f"Fonte: {fonte or ANO_FONTE}.", acima=False, fonte=True)
    doc.add_paragraph()


def lista(doc, itens):
    for it in itens:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        partes = it.split("**")
        for i, parte in enumerate(partes):
            r = p.add_run(parte); r.font.name = FONTE; r.font.size = Pt(12)
            if i % 2 == 1: r.bold = True


def referencias(doc, numero, lista_refs):
    """Seção de Referências (NBR 6023): alinhamento à esquerda, espaço simples."""
    secao(doc, numero, "Referências", nivel=1)
    for ref in sorted(lista_refs):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(12)
        # primeira parte (sobrenome em maiúsculas) já vem formatada no texto
        r = p.add_run(ref); r.font.name = FONTE; r.font.size = Pt(12)
